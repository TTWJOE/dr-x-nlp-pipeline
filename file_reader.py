import os
import fitz  # 
import docx
import pandas as pd

SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.csv', '.xlsx', '.xls', '.xlsm']

def extract_text_from_docx(path):
    doc = docx.Document(path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))  # readable format
    return '\n'.join(full_text)

def extract_text_from_pdf(path):
    doc = fitz.open(path)
    full_text = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        full_text.append(f"[Page {page_num}]\n{text}")
    return '\n'.join(full_text)

def extract_text_from_spreadsheet(path):
    try:
        df_dict = pd.read_excel(path, sheet_name=None, engine='openpyxl')
    except Exception:
        df_dict = pd.read_excel(path, sheet_name=None, engine='xlrd')

    all_text = []
    for sheet, df in df_dict.items():
        all_text.append(f"--- Sheet: {sheet} ---")
        all_text.append(df.to_string(index=False, header=True))  # Flatten table to text
    return "\n".join(all_text)

def extract_text_from_csv(path):
    df = pd.read_csv(path)
    return df.to_string(index=False, header=True)

def extract_text_from_file(path):
    ext = os.path.splitext(path)[-1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file format: {ext}")

    try:
        if ext == '.pdf':
            return extract_text_from_pdf(path)
        elif ext == '.docx':
            return extract_text_from_docx(path)
        elif ext == '.csv':
            return extract_text_from_csv(path)
        elif ext in ['.xlsx', '.xls', '.xlsm']:
            return extract_text_from_spreadsheet(path)
    except Exception as e:
        print(f"Error reading {path}: {e}")
        return ""

from pathlib import Path

if __name__ == "__main__":
    folder = Path("files/")
    for file_path in folder.glob("*"):
        print(f"Reading {file_path.name}...")
        text = extract_text_from_file(str(file_path))
        print(text[:500])  # Print preview